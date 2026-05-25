"""DOCX 导出模块 - 将翻译结果保存为 DOCX 文件"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _is_separator_row(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行（如 |---|---|）"""
    stripped = line.strip().replace('|', '').strip()
    return bool(stripped) and all(c in '-: ' for c in stripped)


def _parse_table_cells(line: str) -> list[str]:
    """解析 Markdown 表格行中的单元格"""
    return [c.strip() for c in line.strip().strip('|').split('|')]


def _add_table_to_doc(doc: Document, table_lines: list[str]):
    """将 Markdown 表格转换为 Word 原生表格"""
    # 过滤掉分隔行
    data_rows = [ln for ln in table_lines if not _is_separator_row(ln)]
    if len(data_rows) < 2:
        return

    # 解析所有行，统一列数
    parsed_rows = [_parse_table_cells(ln) for ln in data_rows]
    max_cols = max(len(r) for r in parsed_rows)

    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)

    # 设置表格样式为带边框的网格
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(parsed_rows):
        for col_idx in range(max_cols):
            cell = row_data[col_idx] if col_idx < len(row_data) else ''
            cell_obj = table.cell(row_idx, col_idx)
            cell_obj.text = ''

            p = cell_obj.paragraphs[0]
            run = p.add_run(cell)

            # 首行作为表头：加粗 + 灰色底色
            if row_idx == 0:
                run.bold = True
                # 设置单元格底色
                shading = cell_obj._tc.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'D9E2F3',
                })
                shading.append(shading_elm)

    # 表格后空行
    doc.add_paragraph()


def _process_markdown_text(doc: Document, text: str):
    """将 Markdown 文本按块解析并写入 DOCX（表格、标题、列表等）"""
    lines = text.strip().split('\n')
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # 检测是否为表格块
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # 至少需包含分隔行才视为表格
            if any(_is_separator_row(ln) for ln in table_lines):
                _add_table_to_doc(doc, table_lines)
            else:
                # 没有分隔行 → 当作普通文本
                for ln in table_lines:
                    doc.add_paragraph(ln)
            continue

        # 水平分隔线
        if re.match(r'^[-*_]{3,}$', line):
            doc.add_paragraph('─' * 40)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+[\.\)]\s', line):
            doc.add_paragraph(line, style='List Number')
            i += 1
            continue

        # 下级标题（##）
        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(13)
            i += 1
            continue

        # 一级标题（#）
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(15)
            i += 1
            continue

        # 无序列表
        if line.startswith('- '):
            doc.add_paragraph(line, style='List Bullet')
            i += 1
            continue

        # 普通段落
        doc.add_paragraph(line)
        i += 1


def generate_translated_docx(
    pdf_path: str,
    translations: dict[int, str],
    total_pages: int,
) -> str:
    """
    根据翻译结果生成 DOCX 文件，保存在原 PDF 同目录。
    返回生成的 DOCX 文件路径。
    """
    doc = Document()

    # 设置默认字号
    style = doc.styles['Normal']
    style.font.size = Pt(11)

    # 标题
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    title = doc.add_heading(f'{base_name} — 中文翻译', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    for page_num in range(1, total_pages + 1):
        text = translations.get(page_num, '')

        # 页码标题
        doc.add_heading(f'第 {page_num} 页', level=2)

        if not text:
            doc.add_paragraph('[未翻译]')
        elif text == '[翻译失败]':
            doc.add_paragraph('[翻译失败]')
        else:
            _process_markdown_text(doc, text)

        doc.add_paragraph()  # 页间空行

    # 输出路径
    output_dir = os.path.dirname(pdf_path) or os.getcwd()
    output_name = f"{base_name}_翻译.docx"
    output_path = os.path.join(output_dir, output_name)

    doc.save(output_path)
    return output_path